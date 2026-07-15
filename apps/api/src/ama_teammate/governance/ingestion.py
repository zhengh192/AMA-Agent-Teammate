from __future__ import annotations

import csv
import io
import re
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from ama_teammate.governance.models import ParsedChunk, SourceLocation

PARSER_VERSION = "phase3-parser-1"
ALLOWED_MEDIA_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".csv": "text/csv",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
}
ACTIVE_ZIP_PATTERNS = ("vbaproject.bin", "externallinks/", "embeddings/")


class UnsafeDocumentError(ValueError):
    pass


def validate_upload(filename: str, media_type: str | None, data: bytes, max_bytes: int) -> str:
    suffix = Path(filename).suffix.lower()
    expected = ALLOWED_MEDIA_TYPES.get(suffix)
    if expected is None:
        raise UnsafeDocumentError("Unsupported file type.")
    if not data or len(data) > max_bytes:
        raise UnsafeDocumentError("File is empty or exceeds the configured size limit.")
    supplied = (media_type or "").split(";", 1)[0].lower()
    compatible = supplied in {"", expected, "application/octet-stream"}
    if suffix in {".txt", ".md", ".markdown"}:
        compatible = compatible or supplied.startswith("text/")
    if not compatible:
        raise UnsafeDocumentError("File extension and media type do not agree.")
    if suffix == ".pdf" and not data.startswith(b"%PDF-"):
        raise UnsafeDocumentError("PDF signature is invalid.")
    if suffix in {".docx", ".xlsx"}:
        _validate_office_zip(suffix, data)
    if suffix in {".csv", ".txt", ".md", ".markdown"}:
        if b"\x00" in data:
            raise UnsafeDocumentError("Text upload contains binary content.")
        try:
            data.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise UnsafeDocumentError("Text uploads must be UTF-8 encoded.") from exc
    return expected


def parse_document(filename: str, data: bytes) -> list[ParsedChunk]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(data)
    if suffix == ".docx":
        return _parse_docx(data)
    if suffix == ".xlsx":
        return _parse_xlsx(data)
    if suffix == ".csv":
        return _parse_csv(data)
    return _parse_text(data.decode("utf-8-sig"), markdown=suffix in {".md", ".markdown"})


def _validate_office_zip(suffix: str, data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            entries = archive.infolist()
            if len(entries) > 5_000:
                raise UnsafeDocumentError("Office archive has too many entries.")
            total = sum(entry.file_size for entry in entries)
            compressed = sum(max(entry.compress_size, 1) for entry in entries)
            if total > 100_000_000 or total / compressed > 100:
                raise UnsafeDocumentError("Office archive exceeds expansion limits.")
            names = [entry.filename.lower() for entry in entries]
            if any(pattern in name for pattern in ACTIVE_ZIP_PATTERNS for name in names):
                raise UnsafeDocumentError("Macros, embedded objects, or external links are disabled.")
            required = "word/document.xml" if suffix == ".docx" else "xl/workbook.xml"
            if required not in names:
                raise UnsafeDocumentError("Office document signature is invalid.")
    except zipfile.BadZipFile as exc:
        raise UnsafeDocumentError("Office document archive is invalid.") from exc


def _parse_pdf(data: bytes) -> list[ParsedChunk]:
    reader = PdfReader(io.BytesIO(data), strict=True)
    if reader.is_encrypted or len(reader.pages) > 250:
        raise UnsafeDocumentError("Encrypted or oversized PDF is not supported.")
    chunks: list[ParsedChunk] = []
    for number, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            chunks.extend(_split(text, SourceLocation(page=number)))
    return chunks


def _parse_docx(data: bytes) -> list[ParsedChunk]:
    document = Document(io.BytesIO(data))
    if len(document.paragraphs) > 20_000:
        raise UnsafeDocumentError("DOCX exceeds paragraph limits.")
    chunks: list[ParsedChunk] = []
    section = "Document"
    buffer: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            if buffer:
                chunks.extend(_split("\n".join(buffer), SourceLocation(section=section)))
                buffer = []
            section = text
        else:
            buffer.append(text)
    if buffer:
        chunks.extend(_split("\n".join(buffer), SourceLocation(section=section)))
    return chunks


def _parse_xlsx(data: bytes) -> list[ParsedChunk]:
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True, keep_links=False)
    try:
        if len(workbook.sheetnames) > 100:
            raise UnsafeDocumentError("XLSX exceeds sheet limits.")
        chunks: list[ParsedChunk] = []
        for worksheet in workbook.worksheets:
            rows: list[str] = []
            start = 1
            for number, values in enumerate(worksheet.iter_rows(values_only=True), 1):
                if number > 5_000:
                    raise UnsafeDocumentError("XLSX exceeds row limits.")
                row = " | ".join("" if value is None else str(value) for value in values).strip()
                if row.strip(" |"):
                    rows.append(row)
                if len(rows) >= 40:
                    chunks.append(
                        ParsedChunk(
                            content="\n".join(rows),
                            location=SourceLocation(
                                sheet=worksheet.title, row_start=start, row_end=number
                            ),
                        )
                    )
                    rows = []
                    start = number + 1
            if rows:
                chunks.append(
                    ParsedChunk(
                        content="\n".join(rows),
                        location=SourceLocation(
                            sheet=worksheet.title,
                            row_start=start,
                            row_end=start + len(rows) - 1,
                        ),
                    )
                )
        return chunks
    finally:
        workbook.close()


def _parse_csv(data: bytes) -> list[ParsedChunk]:
    reader = csv.reader(io.StringIO(data.decode("utf-8-sig")))
    chunks: list[ParsedChunk] = []
    rows: list[str] = []
    start = 1
    for number, values in enumerate(reader, 1):
        if number > 10_000:
            raise UnsafeDocumentError("CSV exceeds row limits.")
        rows.append(" | ".join(_escape_formula(value) for value in values))
        if len(rows) >= 50:
            chunks.append(
                ParsedChunk(
                    content="\n".join(rows),
                    location=SourceLocation(row_start=start, row_end=number),
                )
            )
            rows = []
            start = number + 1
    if rows:
        chunks.append(
            ParsedChunk(
                content="\n".join(rows),
                location=SourceLocation(row_start=start, row_end=start + len(rows) - 1),
            )
        )
    return chunks


def _parse_text(text: str, *, markdown: bool) -> list[ParsedChunk]:
    lines = text.splitlines()
    if len(lines) > 50_000:
        raise UnsafeDocumentError("Text document exceeds line limits.")
    if not markdown:
        return _split(text, SourceLocation(line_start=1, line_end=max(1, len(lines))))
    chunks: list[ParsedChunk] = []
    section = "Document"
    buffer: list[str] = []
    start = 1
    for number, line in enumerate(lines, 1):
        match = re.match(r"^#{1,6}\s+(.+)$", line)
        if match:
            if buffer:
                chunks.extend(
                    _split(
                        "\n".join(buffer),
                        SourceLocation(section=section, line_start=start, line_end=number - 1),
                    )
                )
            section = match.group(1).strip()
            buffer = []
            start = number + 1
        else:
            buffer.append(line)
    if buffer:
        chunks.extend(
            _split(
                "\n".join(buffer),
                SourceLocation(section=section, line_start=start, line_end=len(lines)),
            )
        )
    return chunks


def _split(text: str, location: SourceLocation) -> list[ParsedChunk]:
    clean = text.strip()
    return [
        ParsedChunk(content=clean[index : index + 3_000], location=location)
        for index in range(0, len(clean), 3_000)
        if clean[index : index + 3_000].strip()
    ]


def _escape_formula(value: str) -> str:
    return "'" + value if value.startswith(("=", "+", "-", "@")) else value
